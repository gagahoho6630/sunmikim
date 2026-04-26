"""Extract content from xlsx/docx/pdf as JSON.

Used by the data-mapper agent. Output is intentionally raw and structural;
the mapping decision (which value goes into which slot) is left to the
agent's judgment, not to this script.
"""
import json
import sys
from pathlib import Path


def extract_xlsx(path: Path) -> dict:
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True)
    sheets = []
    for ws in wb.worksheets:
        rows = [list(r) for r in ws.iter_rows(values_only=True)]
        sheets.append({"name": ws.title, "rows": rows})
    return {"sheets": sheets}


def extract_docx(path: Path) -> dict:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = [
        [[c.text for c in row.cells] for row in t.rows] for t in doc.tables
    ]
    return {"paragraphs": paragraphs, "tables": tables}


def extract_pdf(path: Path) -> dict:
    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(str(path)) as pdf:
            return {
                "pages": [
                    {
                        "text": page.extract_text() or "",
                        "tables": page.extract_tables() or [],
                    }
                    for page in pdf.pages
                ]
            }
    except ImportError:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return {
            "pages": [
                {"text": page.extract_text() or "", "tables": []}
                for page in reader.pages
            ]
        }


EXTRACTORS = {
    ".xlsx": extract_xlsx,
    ".xlsm": extract_xlsx,
    ".docx": extract_docx,
    ".pdf": extract_pdf,
}


def main():
    if len(sys.argv) < 2:
        print("usage: extract_input.py <file>", file=sys.stderr)
        sys.exit(2)
    path = Path(sys.argv[1])
    if not path.exists():
        print(f"input not found: {path}", file=sys.stderr)
        sys.exit(2)
    ext = path.suffix.lower()
    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        print(f"unsupported format: {ext}", file=sys.stderr)
        sys.exit(2)
    data = extractor(path)
    json.dump(
        {"format": ext.lstrip("."), "path": str(path), "data": data},
        sys.stdout,
        ensure_ascii=False,
        indent=2,
        default=str,
    )


if __name__ == "__main__":
    main()
